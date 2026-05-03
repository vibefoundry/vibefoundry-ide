"""
FastAPI backend server for VibeFoundry IDE
"""

import os
import sys
import json
import math
import asyncio
import struct
import signal
import time
from pathlib import Path


def _safe_float_or_none(v):
    """Coerce to JSON-safe float or None. NaN/Inf/None all become None (shown blank in UI)."""
    if v is None:
        return None
    try:
        f = float(v)
    except (TypeError, ValueError):
        return None
    if math.isnan(f) or math.isinf(f):
        return None
    return f


def _safe_int(v):
    """Coerce to int. None/NaN become 0. Use for count-style stats that are always integers ≥ 0."""
    if v is None:
        return 0
    try:
        f = float(v)
    except (TypeError, ValueError):
        return 0
    if math.isnan(f) or math.isinf(f):
        return 0
    return int(f)

# Unix-only imports for terminal functionality
if sys.platform != 'win32':
    import pty
    import fcntl
    import termios
    import select
else:
    pty = None
    fcntl = None
    termios = None
    select = None
from typing import Optional
from contextlib import asynccontextmanager

import httpx
import polars as pl
from fastapi import FastAPI, HTTPException, WebSocket, WebSocketDisconnect, UploadFile, File, Form
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, JSONResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel

from vibefoundry.runner import discover_scripts, run_script, setup_project_structure, ScriptResult, stop_all_scripts, list_running_processes, stop_process
from vibefoundry.metadata import generate_metadata
from vibefoundry.watcher import FileWatcher
from vibefoundry.profiler import (
    is_file_massive, get_profile_cache_path, is_profile_valid,
    profile_large_file, read_cached_profile, estimate_filtered_rows,
    apply_column_exclusions,
    _detect_csv_separator, _get_lazy_frame,
)


# Global state
class AppState:
    project_folder: Optional[Path] = None
    watcher: Optional[FileWatcher] = None
    websocket_clients: list[WebSocket] = []
    # Debounce for script change notifications (prevent duplicates)
    last_script_change: dict[str, float] = {}  # path -> timestamp


class DataFrameState:
    """Stream-from-disk DataFrame viewer - only loads rows as needed"""
    def __init__(self):
        self.file_path: Optional[str] = None
        self.file_type: Optional[str] = None  # 'csv' or 'excel'
        self.csv_separator: str = ','
        self.columns: list[str] = []
        self.column_info: dict = {}  # {col: {type, min, max, values}}
        self.total_rows: int = 0
        self.current_filters: dict = {}
        self.current_sort: Optional[dict] = None
        self.row_limit: Optional[int] = None  # Cap rows for large file preview
        # Small cache for filtered row count (avoids re-scanning)
        self._filtered_row_count: Optional[int] = None

    def clear(self):
        """Clear state"""
        print(f"[Memory] Clearing DataFrame state")
        self.file_path = None
        self.file_type = None
        self.csv_separator = ','
        self.columns = []
        self.column_info = {}
        self.total_rows = 0
        self.current_filters = {}
        self.current_sort = None
        self.row_limit = None
        self._filtered_row_count = None

    def _get_lazy_frame(self) -> Optional[pl.LazyFrame]:
        """Get a lazy frame for the file (doesn't load data)"""
        if not self.file_path:
            return None
        file_path = Path(self.file_path)
        if self.file_type == 'csv':
            return pl.scan_csv(file_path, separator=self.csv_separator, infer_schema_length=10000)
        elif self.file_type == 'parquet':
            return pl.scan_parquet(file_path)
        elif self.file_type == 'excel':
            # Excel doesn't support lazy scanning, load eagerly but this is rare
            return pl.read_excel(file_path).lazy()
        return None

    def _apply_filters_sort(self, lf: pl.LazyFrame) -> pl.LazyFrame:
        """Apply current filters and sort to a lazy frame"""
        try:
            schema = lf.collect_schema()
        except Exception:
            schema = None

        # Apply filters
        for column, filter_val in self.current_filters.items():
            if column not in self.columns:
                continue
            if isinstance(filter_val, dict):
                if 'values' in filter_val:
                    # Categorical filter (object form with optional exclude)
                    vals = filter_val.get('values') or []
                    if vals:
                        str_vals = [str(v) for v in vals]
                        lf = lf.filter(pl.col(column).cast(pl.Utf8).is_in(str_vals))
                else:
                    # Numeric range filter
                    if filter_val.get('min') not in (None, '', 'null'):
                        try:
                            min_val = float(filter_val['min'])
                            lf = lf.filter(pl.col(column).cast(pl.Float64, strict=False) >= min_val)
                        except (ValueError, TypeError):
                            pass
                    if filter_val.get('max') not in (None, '', 'null'):
                        try:
                            max_val = float(filter_val['max'])
                            lf = lf.filter(pl.col(column).cast(pl.Float64, strict=False) <= max_val)
                        except (ValueError, TypeError):
                            pass
                lf = apply_column_exclusions(lf, column, filter_val.get('exclude') or [], schema)
            elif isinstance(filter_val, list) and len(filter_val) > 0:
                # Categorical filter with special sentinels for null/blank/zero
                SPECIAL = {'__vf_filter_null__', '__vf_filter_blank__', '__vf_filter_zero__'}
                specials = [v for v in filter_val if v in SPECIAL]
                regular = [v for v in filter_val if v not in SPECIAL]
                predicates = []
                if regular:
                    str_vals = [str(v) for v in regular]
                    predicates.append(pl.col(column).cast(pl.Utf8).is_in(str_vals))
                if '__vf_filter_null__' in specials:
                    predicates.append(pl.col(column).is_null())
                if '__vf_filter_blank__' in specials:
                    predicates.append(pl.col(column).cast(pl.Utf8) == '')
                if '__vf_filter_zero__' in specials:
                    predicates.append(pl.col(column).cast(pl.Float64, strict=False) == 0)
                if predicates:
                    combined = predicates[0]
                    for p in predicates[1:]:
                        combined = combined | p
                    lf = lf.filter(combined)

        # Apply sort
        if self.current_sort and self.current_sort.get('column'):
            sort_col = self.current_sort['column']
            descending = self.current_sort.get('direction', 'asc') != 'asc'
            if sort_col in self.columns:
                lf = lf.sort(sort_col, descending=descending, nulls_last=True)

        return lf

    def get_rows(self, offset: int, limit: int) -> tuple[list[dict], int]:
        """Get rows with current filters/sort applied. Returns (rows, total_filtered_count)"""
        lf = self._get_lazy_frame()
        if lf is None:
            return [], 0

        lf = self._apply_filters_sort(lf)

        # Apply row limit if set (large file preview)
        if self.row_limit is not None:
            lf = lf.slice(0, self.row_limit)

        # Get total count (cached if no filter changes)
        if self._filtered_row_count is None:
            self._filtered_row_count = lf.select(pl.len()).collect().item()

        # Get requested slice
        rows_df = lf.slice(offset, limit).collect()
        rows = rows_df.to_dicts()

        # Null → blank, NaN → "NaN" (visible), Inf → blank (for JSON safety)
        for row in rows:
            for key in row:
                v = row[key]
                if v is None:
                    row[key] = ''
                elif isinstance(v, float):
                    if math.isnan(v):
                        row[key] = 'NaN'
                    elif math.isinf(v):
                        row[key] = ''

        return rows, self._filtered_row_count

    def invalidate_filter_cache(self):
        """Call when filters change"""
        self._filtered_row_count = None


state = AppState()
df_state = DataFrameState()

# Track active profiling task so it can be cancelled
_profiling_task: Optional[asyncio.Task] = None


def _compute_column_info(lf: pl.LazyFrame, columns: list, schema) -> dict:
    """Compute column info in a single optimized pass.
    Batches all numeric stats into one query, then handles categorical columns.
    Returns min/max/nullCount/zeroCount for numeric, values/nullCount/blankCount for categorical."""
    column_info = {}

    # Separate numeric and categorical columns
    numeric_cols = []
    categorical_cols = []
    for col in columns:
        dtype = schema.get(col)
        if dtype is None:
            continue
        if dtype.is_numeric():
            numeric_cols.append(col)
        else:
            categorical_cols.append(col)

    # Batch all numeric column stats in ONE query (single file scan)
    if numeric_cols:
        try:
            float_cols = {c for c in numeric_cols if schema.get(c) in (pl.Float32, pl.Float64)}
            exprs = []
            for col in numeric_cols:
                # fill_nan(None) so NaN is ignored in aggregations (treated as missing)
                clean = pl.col(col).fill_nan(None) if col in float_cols else pl.col(col)
                exprs.extend([
                    clean.min().alias(f'{col}__min'),
                    clean.max().alias(f'{col}__max'),
                    clean.sum().alias(f'{col}__sum'),
                    clean.mean().alias(f'{col}__mean'),
                    clean.median().alias(f'{col}__median'),
                    pl.col(col).count().alias(f'{col}__count'),
                    pl.col(col).is_null().sum().alias(f'{col}__null'),
                    (pl.col(col) == 0).sum().alias(f'{col}__zero'),
                    pl.col(col).drop_nulls().n_unique().alias(f'{col}__unique'),
                ])
                if col in float_cols:
                    exprs.append(pl.col(col).is_nan().sum().alias(f'{col}__nan'))
            stats = lf.select(exprs).collect()

            for col in numeric_cols:
                column_info[col] = {
                    "type": "numeric",
                    "min": _safe_float_or_none(stats[f'{col}__min'][0]),
                    "max": _safe_float_or_none(stats[f'{col}__max'][0]),
                    "sum": _safe_float_or_none(stats[f'{col}__sum'][0]),
                    "mean": _safe_float_or_none(stats[f'{col}__mean'][0]),
                    "median": _safe_float_or_none(stats[f'{col}__median'][0]),
                    "count": _safe_int(stats[f'{col}__count'][0]),
                    "nullCount": _safe_int(stats[f'{col}__null'][0]),
                    "zeroCount": _safe_int(stats[f'{col}__zero'][0]),
                    "uniqueCount": _safe_int(stats[f'{col}__unique'][0]),
                    "nanCount": _safe_int(stats[f'{col}__nan'][0]) if col in float_cols else 0,
                }
        except Exception:
            for col in numeric_cols:
                column_info[col] = {"type": "numeric", "min": None, "max": None, "sum": None, "mean": None, "median": None, "count": 0, "nullCount": 0, "zeroCount": 0, "uniqueCount": 0, "nanCount": 0}

    # Batch categorical stats in ONE query
    if categorical_cols:
        try:
            exprs = []
            for col in categorical_cols:
                exprs.extend([
                    pl.col(col).count().alias(f'{col}__count'),
                    pl.col(col).is_null().sum().alias(f'{col}__null'),
                    (pl.col(col).cast(pl.Utf8) == '').sum().alias(f'{col}__blank'),
                    pl.col(col).drop_nulls().n_unique().alias(f'{col}__unique'),
                ])
            stats = lf.select(exprs).collect()

            # Get unique values for each categorical column (requires separate queries for .unique())
            for col in categorical_cols:
                try:
                    unique_vals = lf.select(
                        pl.col(col).drop_nulls().cast(pl.Utf8).unique()
                    ).collect()[col].to_list()
                    unique_vals = sorted([str(v) for v in unique_vals if v != ''])
                except Exception:
                    unique_vals = []

                column_info[col] = {
                    "type": "categorical",
                    "values": unique_vals,
                    "count": _safe_int(stats[f'{col}__count'][0]),
                    "nullCount": _safe_int(stats[f'{col}__null'][0]),
                    "blankCount": _safe_int(stats[f'{col}__blank'][0]),
                    "uniqueCount": _safe_int(stats[f'{col}__unique'][0]),
                }
        except Exception:
            for col in categorical_cols:
                column_info[col] = {"type": "categorical", "values": [], "count": 0, "nullCount": 0, "blankCount": 0, "uniqueCount": 0}

    return column_info


# Alias for backward compatibility
_compute_full_column_info = _compute_column_info


# Request/Response models
class FolderSelectRequest(BaseModel):
    path: str


class RunScriptsRequest(BaseModel):
    scripts: list[str]


class ScriptResultResponse(BaseModel):
    script_path: str
    success: bool
    stdout: str
    stderr: str
    return_code: int
    error: Optional[str] = None
    timed_out: bool = False
    streamlit_url: Optional[str] = None  # URL if this was a Streamlit app


@asynccontextmanager
async def lifespan(app: FastAPI):
    """Startup and shutdown events"""
    # Check for project folder from environment
    project_path = os.environ.get("VIBEFOUNDRY_PROJECT_PATH")
    if project_path:
        folder = Path(project_path)
        if folder.exists() and folder.is_dir():
            state.project_folder = folder
            setup_project_structure(folder)
            generate_metadata(folder)
            state.watcher = FileWatcher(folder)
            state.watcher.scan_initial_state()

    yield
    # Cleanup
    if state.watcher:
        state.watcher.stop()
    # Stop any running scripts (including Streamlit apps)
    stopped = stop_all_scripts()
    if stopped:
        print(f"[Shutdown] Stopped {stopped} running script(s)")


# Create FastAPI app
app = FastAPI(
    title="VibeFoundry IDE",
    version="0.1.0",
    lifespan=lifespan
)

# CORS for development
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


def get_static_dir() -> Path:
    """Get the path to bundled static files"""
    return Path(__file__).parent / "static"


# API Routes

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "ok", "project_folder": str(state.project_folder) if state.project_folder else None}


class LaunchTerminalRequest(BaseModel):
    path: str
    command: str = None  # Optional command to run after cd (e.g., 'claude', 'codex')


@app.post("/api/terminal/launch")
async def launch_native_terminal(request: LaunchTerminalRequest):
    """Launch a native terminal window, cd into the project, and optionally run a command"""
    import subprocess

    folder_path = Path(request.path)
    if not folder_path.exists():
        raise HTTPException(status_code=400, detail="Folder does not exist")

    if sys.platform == 'darwin':  # macOS
        # Use AppleScript to open Terminal.app with commands
        if request.command:
            script = f'''
            tell application "Terminal"
                activate
                do script "cd \\"{folder_path}\\" && clear && {request.command}"
            end tell
            '''
        else:
            script = f'''
            tell application "Terminal"
                activate
                do script "cd \\"{folder_path}\\" && clear"
            end tell
            '''
        subprocess.run(['osascript', '-e', script], check=True)
        return {"status": "ok", "message": "Terminal launched"}
    elif sys.platform == 'win32':  # Windows
        # Use start command with /d to set working directory
        if request.command:
            subprocess.Popen(
                f'start "" /d "{folder_path}" cmd /k {request.command}',
                shell=True
            )
        else:
            subprocess.Popen(
                f'start "" /d "{folder_path}" cmd',
                shell=True
            )
        return {"status": "ok", "message": "Terminal launched"}
    else:
        raise HTTPException(status_code=400, detail="Native terminal launch not supported on this platform")


@app.post("/api/folder/select")
async def select_folder(request: FolderSelectRequest):
    """Set the project folder and initialize structure"""
    folder_path = Path(request.path)

    if not folder_path.exists():
        raise HTTPException(status_code=400, detail="Folder does not exist")

    if not folder_path.is_dir():
        raise HTTPException(status_code=400, detail="Path is not a directory")

    state.project_folder = folder_path

    # Don't auto-scaffold - user must click Build button
    # Just ensure basic folders exist for watcher
    folders = {
        "input_folder": folder_path / "input_folder",
        "output_folder": folder_path / "output_folder",
        "app_folder": folder_path / "app_folder",
        "scripts_folder": folder_path / "app_folder" / "scripts",
    }

    # Stop existing watcher
    if state.watcher:
        state.watcher.stop()

    # Start new watcher
    # Note: Pass coroutines directly - watcher.py handles thread-safe scheduling
    state.watcher = FileWatcher(
        folder_path,
        on_data_change=notify_data_change,
        on_script_change=notify_script_change,
        on_output_file_change=notify_output_file_change
    )
    await state.watcher.start_async()

    # Generate initial metadata
    generate_metadata(folder_path)

    return {
        "success": True,
        "name": folder_path.name,
        "project_folder": str(folder_path),
        "folders": {k: str(v) for k, v in folders.items()}
    }


@app.get("/api/folder/info")
async def get_folder_info():
    """Get current project folder info"""
    if not state.project_folder:
        return {"project_folder": None}

    return {
        "project_folder": str(state.project_folder),
        "name": state.project_folder.name
    }


@app.post("/api/build")
async def build_project():
    """Build the project structure - creates folders and copies instruction files"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    # Create folder structure
    folders = setup_project_structure(state.project_folder)

    import shutil
    import urllib.request
    templates_dir = Path(__file__).parent / "templates"
    TEMPLATE_BASE_URL = "https://vibefoundry.ai/templates"

    for filename in ("AGENTS.md",):
        dest = state.project_folder / filename
        # Try downloading the latest version from the website
        try:
            url = f"{TEMPLATE_BASE_URL}/{filename}"
            urllib.request.urlretrieve(url, str(dest))
        except Exception:
            # Fallback to bundled copy
            bundled = templates_dir / filename
            if bundled.exists():
                shutil.copy2(bundled, dest)

    # Initialize git repo if not already one
    git_initialized = False
    git_dir = state.project_folder / ".git"
    if not git_dir.exists():
        import subprocess
        try:
            subprocess.run(
                ["git", "init"],
                cwd=str(state.project_folder),
                capture_output=True,
                check=True
            )
            git_initialized = True

            # Create .gitignore with sensible defaults
            gitignore_path = state.project_folder / ".gitignore"
            if not gitignore_path.exists():
                gitignore_path.write_text(
                    "# Python\n"
                    "__pycache__/\n"
                    "*.py[cod]\n"
                    ".venv/\n"
                    "venv/\n"
                    "*.egg-info/\n"
                    "\n"
                    "# Node\n"
                    "node_modules/\n"
                    "\n"
                    "# Environment\n"
                    ".env\n"
                    ".env.local\n"
                    "\n"
                    "# OS\n"
                    ".DS_Store\n"
                    "Thumbs.db\n",
                    encoding="utf-8"
                )
        except (subprocess.CalledProcessError, FileNotFoundError):
            # git not installed or failed - continue without it
            pass

    # Generate metadata now that folders exist
    generate_metadata(state.project_folder)

    # Restart watcher to pick up newly created folders
    if state.watcher:
        state.watcher.stop()
    state.watcher = FileWatcher(
        state.project_folder,
        on_data_change=notify_data_change,
        on_script_change=notify_script_change,
        on_output_file_change=notify_output_file_change
    )
    await state.watcher.start_async()

    return {
        "success": True,
        "folders": {k: str(v) for k, v in folders.items()},
        "agents_md_copied": (state.project_folder / "AGENTS.md").exists(),
        "git_initialized": git_initialized
    }


@app.get("/api/scripts")
async def list_scripts():
    """List available scripts"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    scripts_folder = state.project_folder / "app_folder" / "scripts"
    scripts = discover_scripts(scripts_folder)

    return {
        "scripts": [
            {
                "path": str(s),
                "relative_path": str(s.relative_to(scripts_folder)),
                "name": s.name
            }
            for s in scripts
        ]
    }


@app.post("/api/scripts/run")
async def run_scripts(request: RunScriptsRequest):
    """Run selected scripts"""
    import asyncio

    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    results: list[ScriptResultResponse] = []

    for script_path in request.scripts:
        # Run in thread pool so server stays responsive (allows stop requests)
        result = await asyncio.to_thread(run_script, Path(script_path), state.project_folder)
        results.append(ScriptResultResponse(
            script_path=result.script_path,
            success=result.success,
            stdout=result.stdout,
            stderr=result.stderr,
            return_code=result.return_code,
            error=result.error,
            timed_out=result.timed_out,
            streamlit_url=result.streamlit_url
        ))

    # Regenerate metadata after running scripts (skip for .sh/.bat since they are long-running apps)
    ran_only_launchers = all(
        Path(s).suffix.lower() in (".sh", ".bat") for s in request.scripts
    )
    if not ran_only_launchers:
        generate_metadata(state.project_folder)

    return {"results": [r.model_dump() for r in results]}


class RunExternalRequest(BaseModel):
    scriptPath: str


@app.post("/api/scripts/run-external")
async def run_script_external(request: RunExternalRequest):
    """Launch a script in the system's external terminal."""
    import subprocess

    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    script_path = Path(request.scriptPath)
    if not script_path.exists():
        raise HTTPException(status_code=404, detail="Script not found")

    cwd = str(state.project_folder)
    script = str(script_path)

    if sys.platform == 'darwin':
        # macOS: open Terminal.app and run the script
        apple_script = f'''tell application "Terminal"
    activate
    do script "cd {cwd} && python \\"{script}\\""
end tell'''
        subprocess.Popen(['osascript', '-e', apple_script])
    elif sys.platform == 'win32':
        # Windows: open cmd and run the script
        subprocess.Popen(
            f'start cmd /k "cd /d {cwd} && python \\"{script}\\""',
            shell=True
        )
    else:
        # Linux: try common terminal emulators
        for term_cmd in [
            ['gnome-terminal', '--', 'bash', '-c', f'cd {cwd} && python "{script}"; exec bash'],
            ['xterm', '-e', f'cd {cwd} && python "{script}"; bash'],
        ]:
            try:
                subprocess.Popen(term_cmd)
                break
            except FileNotFoundError:
                continue

    return {"success": True, "scriptPath": request.scriptPath}


@app.post("/api/scripts/stop")
async def stop_scripts():
    """Stop all currently running scripts"""
    stopped = stop_all_scripts()
    print(f"[Scripts] Stopped {stopped} running script(s)")
    return {"success": True, "stopped": stopped}


@app.get("/api/processes")
async def get_running_processes():
    """List all currently running script processes"""
    processes = list_running_processes()
    return {"processes": processes}


class StopProcessRequest(BaseModel):
    pid: int


@app.post("/api/processes/stop")
async def stop_single_process(request: StopProcessRequest):
    """Stop a specific process by PID"""
    success = stop_process(request.pid)
    if success:
        print(f"[Processes] Stopped process {request.pid}")
        return {"success": True, "pid": request.pid}
    else:
        return {"success": False, "error": f"Process {request.pid} not found or could not be stopped"}


@app.post("/api/metadata/generate")
async def regenerate_metadata():
    """Force metadata regeneration"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    input_meta, output_meta = await asyncio.to_thread(generate_metadata, state.project_folder)

    return {
        "success": True,
        "input_metadata": input_meta,
        "output_metadata": output_meta
    }


class PipInstallRequest(BaseModel):
    package: str


@app.post("/api/pip/install")
async def pip_install(request: PipInstallRequest):
    """Install a Python package using pip"""
    import subprocess
    import sys

    # Sanitize package name - only allow alphanumeric, hyphens, underscores, brackets
    package = request.package.strip()
    if not package or not all(c.isalnum() or c in '-_[],' for c in package):
        raise HTTPException(status_code=400, detail="Invalid package name")

    try:
        # Run pip install
        result = subprocess.run(
            [sys.executable, "-m", "pip", "install", package],
            capture_output=True,
            text=True,
            timeout=120  # 2 minute timeout
        )

        return {
            "success": result.returncode == 0,
            "package": package,
            "stdout": result.stdout,
            "stderr": result.stderr,
            "return_code": result.returncode
        }
    except subprocess.TimeoutExpired:
        return {
            "success": False,
            "package": package,
            "stdout": "",
            "stderr": "Installation timed out",
            "return_code": -1
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to install package: {str(e)}")


@app.get("/api/watch/check")
async def check_for_changes():
    """Manually check for file changes"""
    if not state.watcher:
        return {"changes": False}

    input_changes, output_changes, script_changes = state.watcher.check_once()

    has_changes = bool(input_changes or output_changes or script_changes)

    if input_changes or output_changes:
        await asyncio.to_thread(generate_metadata, state.project_folder)

    return {
        "changes": has_changes,
        "input_changes": [{"path": c.path, "type": c.change_type} for c in input_changes],
        "output_changes": [{"path": c.path, "type": c.change_type} for c in output_changes],
        "script_changes": [{"path": c.path, "type": c.change_type} for c in script_changes]
    }


# Filesystem browsing endpoints

@app.get("/api/fs/home")
async def get_home_directory():
    """Get user's home directory"""
    return {"path": str(Path.home())}


@app.get("/api/fs/list")
async def list_directory(path: str = ""):
    """List directories at a given path (for folder picker)"""
    if not path:
        path = str(Path.home())

    target = Path(path)

    if not target.exists():
        raise HTTPException(status_code=404, detail="Path does not exist")

    if not target.is_dir():
        raise HTTPException(status_code=400, detail="Path is not a directory")

    folders = []
    try:
        for item in sorted(target.iterdir()):
            # Only show directories, skip hidden files
            if item.is_dir() and not item.name.startswith('.'):
                folders.append({
                    "name": item.name,
                    "path": str(item)
                })
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied")

    return {
        "current": str(target),
        "parent": str(target.parent) if target.parent != target else None,
        "folders": folders
    }


class MkdirRequest(BaseModel):
    path: str
    name: str


@app.post("/api/fs/mkdir")
async def create_directory(request: MkdirRequest):
    """Create a new directory"""
    parent = Path(request.path)

    # If path is relative, make it relative to project folder
    if not parent.is_absolute() and state.project_folder:
        parent = state.project_folder / request.path

    if not parent.exists():
        raise HTTPException(status_code=404, detail=f"Parent path does not exist: {parent}")

    if not parent.is_dir():
        raise HTTPException(status_code=400, detail=f"Parent path is not a directory: {parent}")

    # Sanitize folder name - no path traversal
    name = request.name.strip()
    if not name or '/' in name or '\\' in name or name.startswith('.'):
        raise HTTPException(status_code=400, detail="Invalid folder name")

    new_folder = parent / name

    if new_folder.exists():
        raise HTTPException(status_code=409, detail=f"Folder already exists: {new_folder}")

    try:
        new_folder.mkdir(parents=False, exist_ok=False)
        return {"success": True, "path": str(new_folder)}
    except PermissionError:
        raise HTTPException(status_code=403, detail="Permission denied")
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to create folder: {str(e)}")


# Directories to skip when building the file tree (heavy/irrelevant for the IDE)
TREE_BLACKLIST = {
    'node_modules', '__pycache__', '.next', 'build', 'dist',
    '.cache', '.parcel-cache', '.turbo', 'coverage',
    'env', 'venv', '.venv',
}


def build_file_tree(path: Path, base_path: Path, deleted_files: list = None, in_app_folder: bool = False) -> dict:
    """Build a file tree recursively"""
    if deleted_files is None:
        deleted_files = []

    rel_path = path.relative_to(base_path).as_posix()
    is_file = path.is_file()
    node = {
        "name": path.name,
        "path": "" if rel_path == "." else rel_path,
        "isDirectory": not is_file,
        "extension": path.suffix if is_file else None,
        "lastModified": path.stat().st_mtime if is_file else None,
    }

    if path.is_dir():
        children = []
        # Check if we're entering app_folder
        entering_app_folder = in_app_folder or path.name == "app_folder"
        try:
            for item in sorted(path.iterdir()):
                # Skip hidden files
                if item.name.startswith('.'):
                    continue
                # Skip blacklisted directories
                if item.is_dir() and item.name in TREE_BLACKLIST:
                    continue

                children.append(build_file_tree(item, base_path, deleted_files, entering_app_folder))
        except PermissionError:
            pass
        node["children"] = children

    return node


@app.get("/api/files/tree")
async def get_file_tree():
    """Get the complete file tree for the project"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    deleted_files = []
    tree = await asyncio.to_thread(build_file_tree, state.project_folder, state.project_folder, deleted_files)
    return {"tree": tree, "deletedFiles": deleted_files}


@app.get("/api/files/read")
async def read_file(path: str, sheet: Optional[str] = None):
    """Read a file's content - streams from disk, doesn't hold data in memory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    # Clear any previous DataFrame state
    if df_state.file_path is not None:
        df_state.clear()

    file_path = state.project_folder / path
    print(f"[File Read] Loading: {path}")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    if not file_path.is_file():
        raise HTTPException(status_code=400, detail="Path is not a file")

    # Security check - ensure path is within project folder
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    # Determine file type and read accordingly
    ext = file_path.suffix.lower()
    binary_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.webp', '.pdf', '.zip', '.tar', '.gz'}
    dataframe_extensions = {'.csv', '.xlsx', '.xls', '.parquet', '.geoparquet'}

    if ext in dataframe_extensions:
        print(f"[File Read] Parsing dataframe: {path}")
        temp_file_path = None  # Track temp files for cleanup
        excel_sheet_names = None  # Track Excel sheet names
        excel_active_sheet = None

        # ── Try metadata first for column names & row count ──
        # This avoids scanning the file at all for basic info
        meta_columns = None
        meta_rows = None
        try:
            from vibefoundry.metadata import _metadata_cache
            cache_key = str(file_path)
            mtime = file_path.stat().st_mtime
            if cache_key in _metadata_cache:
                cached_mtime, cached_rows, cached_cols, cached_info = _metadata_cache[cache_key]
                if cached_mtime == mtime:
                    meta_columns = cached_cols
                    meta_rows = cached_rows
                    print(f"[File Read] Using cached metadata: {len(meta_columns)} cols, {meta_rows} rows")
        except Exception:
            pass

        # Also check the metadata txt files for row/column info
        if meta_columns is None:
            try:
                for meta_name in ("input_metadata.txt", "output_metadata.txt"):
                    meta_path = state.project_folder / "app_folder" / "meta_data" / meta_name
                    if meta_path.exists():
                        meta_text = meta_path.read_text(encoding="utf-8")
                        # Look for this file's entry in metadata
                        fname = file_path.name
                        if f"File: {fname}" in meta_text or str(file_path) in meta_text:
                            import re
                            # Extract row count
                            for line in meta_text.split('\n'):
                                if fname in line or str(file_path.relative_to(state.project_folder)) in line:
                                    # Found the file section, scan next lines
                                    pass
                            # Parse rows from metadata
                            sections = meta_text.split("File: ")
                            for section in sections:
                                if fname in section or str(file_path) in section:
                                    rows_match = re.search(r'Rows:\s*(\d+)', section)
                                    if rows_match:
                                        meta_rows = int(rows_match.group(1))
                                    cols_match = re.search(r'Columns\s*\((\d+)\):', section)
                                    if cols_match:
                                        # Extract column names from the "    - colname" lines
                                        col_lines = re.findall(r'^\s+-\s+(.+?)(?:\s+\[.*\])?$', section, re.MULTILINE)
                                        if col_lines:
                                            meta_columns = col_lines
                                    break
                            if meta_columns:
                                print(f"[File Read] Using metadata txt: {len(meta_columns)} cols, {meta_rows} rows")
                                break
            except Exception:
                pass

        try:
            if ext == '.csv':
                try:
                    # Read raw bytes to detect line endings and separator
                    with open(file_path, 'rb') as f:
                        sample = f.read(4096)

                    # Detect line ending style
                    has_crlf = b'\r\n' in sample
                    has_lf = b'\n' in sample
                    has_cr = b'\r' in sample

                    # Detect separator from first line
                    if has_crlf:
                        first_line = sample.split(b'\r\n')[0].decode('utf-8', errors='replace')
                    elif has_lf:
                        first_line = sample.split(b'\n')[0].decode('utf-8', errors='replace')
                    elif has_cr:
                        first_line = sample.split(b'\r')[0].decode('utf-8', errors='replace')
                    else:
                        first_line = sample.decode('utf-8', errors='replace')

                    # Detect separator
                    if '\t' in first_line:
                        separator = '\t'
                    elif ';' in first_line:
                        separator = ';'
                    else:
                        separator = ','

                    # Handle old Mac CR-only line endings - need temp file for streaming
                    needs_cr_conversion = has_cr and not has_lf and not has_crlf
                    actual_file_path = file_path

                    if needs_cr_conversion:
                        import tempfile
                        with open(file_path, 'rb') as f:
                            content = f.read()
                        content = content.replace(b'\r', b'\n')
                        tf = tempfile.NamedTemporaryFile(mode='wb', suffix='.csv', delete=False)
                        tf.write(content)
                        tf.close()
                        temp_file_path = tf.name
                        actual_file_path = Path(temp_file_path)
                        del content

                    # Store CSV file info for streaming
                    df_state.clear()
                    df_state.file_path = Path(actual_file_path).as_posix()
                    df_state.csv_separator = separator
                    df_state.file_type = 'csv'

                    # Get schema from a quick scan (only infer from first rows)
                    lf = pl.scan_csv(actual_file_path, separator=separator, infer_schema_length=10000)
                    df_state.columns = lf.collect_schema().names()
                    schema = lf.collect_schema()

                    # CSVs never go through massive file filtering — they get auto-converted
                    # to Parquet on upload if >50MB. Just load normally.
                    # Use metadata row count if available, otherwise count
                    if meta_rows is not None:
                        df_state.total_rows = meta_rows
                    else:
                        df_state.total_rows = lf.select(pl.len()).collect().item()

                except Exception as csv_err:
                    return {"type": "error", "message": f"Could not read CSV file: {csv_err}", "filename": file_path.name}

            elif ext in {'.parquet', '.geoparquet'}:
                try:
                    df_state.clear()
                    df_state.file_path = file_path.as_posix()
                    df_state.file_type = 'parquet'
                    df_state.csv_separator = ','

                    # Get row count cheaply from parquet metadata if not cached
                    parquet_rows = meta_rows
                    if parquet_rows is None:
                        try:
                            import pyarrow.parquet as pq
                            parquet_rows = pq.ParquetFile(file_path).metadata.num_rows
                        except Exception:
                            parquet_rows = 0

                    # For massive files, use pyarrow metadata to get columns/rows
                    # without scanning the data — avoids OOM on huge files
                    if is_file_massive(file_path, total_rows=parquet_rows):
                        try:
                            import pyarrow.parquet as pq
                            pf = pq.ParquetFile(file_path)
                            arrow_schema = pf.schema_arrow
                            df_state.columns = [f.name for f in arrow_schema]
                            df_state.total_rows = pf.metadata.num_rows
                            schema = {}
                            for field in arrow_schema:
                                schema[field.name] = str(field.type)
                            file_size = file_path.stat().st_size
                            profile_path = get_profile_cache_path(state.project_folder, file_path)
                            has_valid_profile = is_profile_valid(profile_path, file_path)
                            col_dtypes = {col: schema.get(col, "Unknown") for col in df_state.columns}
                            print(f"[File Read] MASSIVE file detected: {file_path.name} ({file_size / 1024 / 1024:.0f} MB). Profile valid: {has_valid_profile}")
                            return {
                                "type": "massive_file",
                                "filename": file_path.name,
                                "filePath": path,
                                "fileSize": file_size,
                                "columns": df_state.columns,
                                "totalRows": df_state.total_rows,
                                "hasProfile": has_valid_profile,
                                "columnDtypes": col_dtypes,
                            }
                        except Exception as pq_err:
                            return {"type": "error", "message": f"Could not read massive Parquet file metadata: {pq_err}", "filename": file_path.name}

                    try:
                        lf = pl.scan_parquet(file_path)
                        df_state.columns = lf.collect_schema().names()
                        schema = lf.collect_schema()
                        # Use metadata row count if available
                        if meta_rows is not None:
                            df_state.total_rows = meta_rows
                        else:
                            df_state.total_rows = lf.select(pl.len()).collect().item()
                    except Exception:
                        # Fallback to pyarrow (e.g. geoparquet with geometry columns)
                        try:
                            import pyarrow.parquet as pq

                            parquet_file = pq.ParquetFile(file_path)
                            arrow_schema = parquet_file.schema_arrow

                            valid_columns = []
                            for field in arrow_schema:
                                if hasattr(field.type, 'extension_name') and 'geo' in str(field.type.extension_name).lower():
                                    continue
                                valid_columns.append(field.name)

                            if not valid_columns:
                                return {"type": "error", "message": "GeoParquet file contains only geometry columns.", "filename": file_path.name}

                            table = parquet_file.read(columns=valid_columns)
                            temp_df = pl.from_arrow(table)

                            df_state.columns = temp_df.columns
                            schema = temp_df.schema
                            df_state.total_rows = meta_rows if meta_rows is not None else len(temp_df)
                            lf = temp_df.lazy()
                            del temp_df
                        except Exception as pyarrow_err:
                            return {"type": "error", "message": f"Could not read Parquet file: {pyarrow_err}", "filename": file_path.name}

                except Exception as parquet_err:
                    return {"type": "error", "message": f"Could not read Parquet file: {parquet_err}", "filename": file_path.name}

            else:
                # Excel (.xlsx, .xls)
                try:
                    from openpyxl import load_workbook
                    df_state.clear()
                    df_state.file_path = file_path.as_posix()
                    df_state.file_type = 'excel'
                    df_state.csv_separator = ','

                    # Get sheet names
                    wb = load_workbook(file_path, read_only=True)
                    excel_sheet_names = wb.sheetnames
                    wb.close()

                    # Read the requested sheet (or first sheet by default)
                    excel_active_sheet = sheet if sheet and sheet in excel_sheet_names else excel_sheet_names[0]
                    target_sheet = excel_active_sheet
                    temp_df = pl.read_excel(file_path, sheet_name=target_sheet)
                    df_state.columns = temp_df.columns
                    schema = temp_df.schema
                    df_state.total_rows = len(temp_df)
                    lf = temp_df.lazy()
                    del temp_df
                except Exception as excel_err:
                    return {"type": "error", "message": f"Could not read Excel file: {excel_err}. Make sure 'openpyxl' is installed (pip install openpyxl).", "filename": file_path.name}

            # Compute column info from the full dataset (Polars lazy scan — efficient)
            try:
                column_info = _compute_full_column_info(lf, df_state.columns, schema)
            except Exception:
                column_info = {col: {"type": "categorical", "values": [], "count": 0, "nullCount": 0, "blankCount": 0, "uniqueCount": 0} for col in df_state.columns}

            df_state.column_info = column_info

            # Get first chunk for initial preview
            CHUNK_SIZE = 200
            first_chunk, total_rows = df_state.get_rows(0, CHUNK_SIZE)

            print(f"[File Read] Fast preview: {df_state.total_rows} total rows, showing {len(first_chunk)}")

            result = {
                "type": "dataframe",
                "filePath": path,
                "columns": df_state.columns,
                "columnInfo": column_info,
                "data": first_chunk,
                "totalRows": df_state.total_rows,
                "offset": 0,
                "limit": CHUNK_SIZE,
                "filename": file_path.name
            }

            # Include sheet info for Excel files
            if excel_sheet_names:
                result["sheetNames"] = excel_sheet_names
                result["activeSheet"] = excel_active_sheet

            return result
        except Exception as e:
            return {"type": "error", "message": f"Unexpected error reading file: {e}", "filename": file_path.name}
        finally:
            # Clean up temp files
            if temp_file_path:
                try:
                    os.remove(temp_file_path)
                except OSError:
                    pass

    elif ext in binary_extensions:
        # Images - return metadata only, frontend uses /api/image endpoint for fast direct loading
        image_extensions = {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.ico', '.webp'}
        if ext in image_extensions:
            return {"type": "image", "path": path, "filename": file_path.name, "extension": ext}
        # PDF files - metadata only, frontend uses /api/pdf endpoint for direct streaming
        if ext == '.pdf':
            return {"type": "pdf", "path": path, "filename": file_path.name}
        # Other binary files - still use base64
        import base64
        content = base64.b64encode(file_path.read_bytes()).decode('utf-8')
        return {"content": content, "encoding": "base64", "filename": file_path.name}
    elif ext == '.json':
        # JSON files - parse and return structured data
        try:
            import json
            content = file_path.read_text(encoding='utf-8')
            data = json.loads(content)
            return {"type": "json", "data": data, "filename": file_path.name}
        except (json.JSONDecodeError, UnicodeDecodeError) as e:
            return {"type": "error", "message": f"Failed to parse JSON: {str(e)}", "filename": file_path.name}
    elif ext == '.docx':
        # Word documents - parse with python-docx
        try:
            from docx import Document
            doc = Document(file_path)
            paragraphs = []
            for para in doc.paragraphs:
                style = para.style.name if para.style else ""
                text = para.text
                if text.strip():
                    paragraphs.append({"text": text, "style": style})

            # Also extract tables
            tables = []
            for table in doc.tables:
                rows = []
                for row in table.rows:
                    rows.append([cell.text for cell in row.cells])
                if rows:
                    tables.append(rows)

            return {
                "type": "docx",
                "paragraphs": paragraphs,
                "tables": tables,
                "filename": file_path.name
            }
        except ImportError:
            return {"type": "error", "message": "Install python-docx to preview Word files: pip install python-docx", "filename": file_path.name}
        except Exception as e:
            return {"type": "error", "message": f"Could not read Word document: {e}", "filename": file_path.name}
    elif ext == '.doc':
        return {"type": "unknown", "message": "Legacy .doc format is not supported. Save as .docx to preview.", "filename": file_path.name}
    else:
        try:
            content = file_path.read_text(encoding='utf-8')
            return {"type": "text", "content": content, "encoding": "utf-8", "filename": file_path.name}
        except UnicodeDecodeError:
            import base64
            content = base64.b64encode(file_path.read_bytes()).decode('utf-8')
            return {"content": content, "encoding": "base64", "filename": file_path.name}


@app.get("/api/image")
async def get_image(path: str):
    """Serve image files directly as binary for fast loading"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / path

    # Security check
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="Image not found")

    # Map extensions to media types
    ext = file_path.suffix.lower()
    media_types = {
        '.png': 'image/png',
        '.jpg': 'image/jpeg',
        '.jpeg': 'image/jpeg',
        '.gif': 'image/gif',
        '.bmp': 'image/bmp',
        '.ico': 'image/x-icon',
        '.webp': 'image/webp',
        '.svg': 'image/svg+xml',
    }

    media_type = media_types.get(ext, 'application/octet-stream')
    return FileResponse(file_path, media_type=media_type)


@app.get("/api/pdf")
async def get_pdf(path: str):
    """Serve PDF files directly with application/pdf media type for inline iframe rendering."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / path

    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists() or file_path.suffix.lower() != ".pdf":
        raise HTTPException(status_code=404, detail="PDF not found")

    return FileResponse(
        file_path,
        media_type="application/pdf",
        headers={"Content-Disposition": f'inline; filename="{file_path.name}"'},
    )


class WriteFileRequest(BaseModel):
    path: str
    content: str


@app.post("/api/files/write")
async def write_file(request: WriteFileRequest):
    """Write content to a file"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.path

    # Security check - ensure path is within project folder
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    # Create parent directories if needed
    file_path.parent.mkdir(parents=True, exist_ok=True)

    file_path.write_text(request.content, encoding='utf-8')

    return {"success": True, "path": request.path}


UPLOAD_CHUNK_SIZE = 8 * 1024 * 1024  # 8MB


@app.post("/api/files/upload")
async def upload_file(
    file: UploadFile = File(...),
    folder: str = Form(...)
):
    """Upload a binary file to a folder, streaming to disk in chunks."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    # Build target path
    target_folder = state.project_folder / folder
    target_path = target_folder / file.filename

    # Security check - ensure path is within project folder
    try:
        target_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    # Create parent directories if needed
    target_folder.mkdir(parents=True, exist_ok=True)

    # Stream file to disk in chunks to avoid loading entire file into memory
    with open(target_path, "wb") as f:
        while True:
            chunk = await file.read(UPLOAD_CHUNK_SIZE)
            if not chunk:
                break
            f.write(chunk)

    result_path = f"{folder}/{target_path.name}"
    return {"success": True, "path": result_path, "converted": False}


class ConvertToParquetRequest(BaseModel):
    path: str
    deleteOriginal: bool = True


@app.post("/api/files/convert-to-parquet")
async def convert_to_parquet(request: ConvertToParquetRequest):
    """Convert a CSV or Excel file to Parquet alongside it. User-triggered only."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    source_path = state.project_folder / request.path
    try:
        source_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not source_path.exists() or not source_path.is_file():
        raise HTTPException(status_code=404, detail="File not found")

    ext = source_path.suffix.lower()
    if ext not in {".csv", ".xlsx", ".xls"}:
        raise HTTPException(status_code=400, detail="Only .csv, .xlsx, and .xls can be converted")

    parquet_path = source_path.with_suffix(".parquet")
    if parquet_path.exists():
        raise HTTPException(status_code=409, detail=f"{parquet_path.name} already exists")

    try:
        if ext == ".csv":
            print(f"[Convert] CSV → Parquet: {source_path.name}")
            pl.scan_csv(
                str(source_path),
                infer_schema_length=10000,
                null_values=["null", "NULL", "None", ""],
            ).sink_parquet(str(parquet_path))
        else:
            print(f"[Convert] Excel → Parquet: {source_path.name}")
            from openpyxl import load_workbook
            wb = load_workbook(source_path, read_only=True)
            sheet_name = wb.sheetnames[0]
            wb.close()
            pl.read_excel(source_path, sheet_name=sheet_name).write_parquet(str(parquet_path))
    except Exception as e:
        if parquet_path.exists():
            parquet_path.unlink()
        raise HTTPException(status_code=500, detail=f"Conversion failed: {e}")

    if request.deleteOriginal:
        try:
            source_path.unlink()
        except Exception as e:
            print(f"[Convert] Could not delete original after convert: {e}")

    rel_parent = Path(request.path).parent.as_posix()
    result_path = f"{rel_parent}/{parquet_path.name}" if rel_parent not in ("", ".") else parquet_path.name
    return {"success": True, "path": result_path}


class DeleteFileRequest(BaseModel):
    path: str
    isDirectory: bool = False


@app.post("/api/files/delete")
async def delete_file(request: DeleteFileRequest):
    """Delete a file or directory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.path

    # Security check - ensure path is within project folder
    try:
        file_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    import shutil

    # Collect files to clean up profiles for (before deletion)
    files_to_clean = []
    try:
        if request.isDirectory:
            for root, _dirs, fnames in os.walk(file_path):
                for fname in fnames:
                    files_to_clean.append(Path(os.path.join(root, fname)))
            shutil.rmtree(file_path)
        else:
            files_to_clean.append(file_path)
            file_path.unlink()
    except PermissionError:
        raise HTTPException(status_code=409, detail="Your File Is Still Open! Close It Before Deleting")

    # Remove any cached profile files for the deleted files
    for f in files_to_clean:
        profile_path = get_profile_cache_path(state.project_folder, f)
        if profile_path.exists():
            profile_path.unlink()
        meta_json = profile_path.with_suffix(".meta.json")
        if meta_json.exists():
            meta_json.unlink()

    # Regenerate metadata so profile files reflect the deletion
    generate_metadata(state.project_folder)

    return {"success": True, "path": request.path}


class RenameRequest(BaseModel):
    oldPath: str
    newName: str


@app.post("/api/files/rename")
async def rename_file(request: RenameRequest):
    """Rename a file or directory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    old_path = Path(request.oldPath)
    if not old_path.is_absolute():
        old_path = state.project_folder / request.oldPath

    # Security check
    try:
        old_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not old_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    new_path = old_path.parent / request.newName

    # Check if new path already exists
    if new_path.exists():
        raise HTTPException(status_code=400, detail="A file with that name already exists")

    import shutil
    shutil.move(str(old_path), str(new_path))

    # Regenerate metadata so profile files reflect the rename
    generate_metadata(state.project_folder)

    return {"success": True, "oldPath": str(old_path), "newPath": str(new_path)}


class MoveRequest(BaseModel):
    sourcePath: str
    destPath: str


@app.post("/api/files/move")
async def move_file(request: MoveRequest):
    """Move a file or directory"""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    source_path = Path(request.sourcePath)
    dest_path = Path(request.destPath)

    if not source_path.is_absolute():
        source_path = state.project_folder / request.sourcePath
    if not dest_path.is_absolute():
        dest_path = state.project_folder / request.destPath

    # Security check
    try:
        source_path.resolve().relative_to(state.project_folder.resolve())
        dest_path.resolve().relative_to(state.project_folder.resolve())
    except ValueError:
        raise HTTPException(status_code=403, detail="Access denied")

    if not source_path.exists():
        raise HTTPException(status_code=404, detail="Source file not found")

    # Ensure destination directory exists
    dest_path.parent.mkdir(parents=True, exist_ok=True)

    import shutil
    shutil.move(str(source_path), str(dest_path))

    # Regenerate metadata so profile files reflect the move
    generate_metadata(state.project_folder)

    return {"success": True, "sourcePath": str(source_path), "destPath": str(dest_path)}


# DataFrame streaming endpoints

class DataFrameQueryRequest(BaseModel):
    filePath: str
    filters: dict = {}
    sort: Optional[dict] = None  # {column: str, direction: "asc"|"desc"}


@app.get("/api/dataframe/rows")
async def get_dataframe_rows(
    filePath: str,
    offset: int = 0,
    limit: int = 200
):
    """Get paginated rows - streams from disk, doesn't hold full file in memory"""
    if df_state.file_path is None:
        raise HTTPException(status_code=400, detail="No DataFrame loaded. Read a file first.")

    # Stream rows from disk
    rows, total_rows = df_state.get_rows(offset, limit)

    return {
        "data": rows,
        "offset": offset,
        "limit": limit,
        "totalRows": total_rows
    }


@app.post("/api/dataframe/query")
async def query_dataframe(request: DataFrameQueryRequest):
    """Apply filters and/or sort to the DataFrame - streams from disk"""
    if df_state.file_path is None:
        raise HTTPException(status_code=400, detail="DataFrame not loaded. Read the file first.")

    # Check if the requested file matches the loaded file (compare by filename since paths may differ)
    loaded_filename = Path(df_state.file_path).name
    requested_filename = Path(request.filePath).name
    if loaded_filename != requested_filename:
        raise HTTPException(status_code=400, detail=f"Different file loaded. Expected {requested_filename}, got {loaded_filename}")

    # Update filters and sort on state
    df_state.current_filters = request.filters
    df_state.current_sort = request.sort
    df_state.invalidate_filter_cache()  # Force recount

    # Get first chunk using streaming
    CHUNK_SIZE = 200
    rows, total_rows = df_state.get_rows(0, CHUNK_SIZE)

    # Compute cascading columnInfo from filtered data
    # For efficiency, we sample a limited number of rows for column stats
    cascading_column_info = await _compute_cascading_column_info()

    return {
        "data": rows,
        "totalRows": total_rows,
        "offset": 0,
        "limit": CHUNK_SIZE,
        "appliedFilters": request.filters,
        "appliedSort": request.sort,
        "columnInfo": cascading_column_info
    }


async def _compute_cascading_column_info() -> dict:
    """Compute column info from filtered data using optimized batched queries."""
    if df_state.file_path is None:
        return {}

    lf = df_state._get_lazy_frame()
    if lf is None:
        return {}

    lf = df_state._apply_filters_sort(lf)
    schema = lf.collect_schema()

    return _compute_column_info(lf, df_state.columns, schema)


@app.post("/api/dataframe/clear")
async def clear_dataframe():
    """Clear the DataFrame from memory"""
    df_state.clear()
    return {"success": True}


# --- Large file profiling endpoints ---

class ProfileRequest(BaseModel):
    filePath: str

class EstimateRequest(BaseModel):
    filePath: str
    filters: dict = {}

class FilteredPreviewRequest(BaseModel):
    filePath: str
    filters: dict = {}
    rowLimit: Optional[int] = None


@app.post("/api/dataframe/profile")
async def start_profile(request: ProfileRequest):
    """Start profiling a massive file. Progress is sent via WebSocket."""
    global _profiling_task

    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.filePath
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    ext = file_path.suffix.lower()
    if ext == ".csv":
        file_type = "csv"
    elif ext in {".parquet", ".geoparquet"}:
        file_type = "parquet"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type for profiling")

    # Check if profile already exists and is valid
    profile_path = get_profile_cache_path(state.project_folder, file_path)
    if is_profile_valid(profile_path, file_path):
        profile = read_cached_profile(profile_path)
        return {"status": "complete", "profile": profile}

    # Cancel any existing profiling task
    if _profiling_task and not _profiling_task.done():
        _profiling_task.cancel()

    async def _run_profiling():
        """Run profiling in a thread and push progress via WebSocket."""
        last_progress = [0]

        def on_progress(done: int, total: int):
            last_progress[0] = done
            # We'll send progress from the async wrapper below
            pass

        # Run the CPU-heavy profiling in a thread
        result = await asyncio.to_thread(
            profile_large_file, file_path, file_type, state.project_folder, on_progress
        )

        # Send completion via WebSocket
        msg = json.dumps({"type": "profile_complete", "filePath": request.filePath, "profile": result})
        disconnected = []
        for client in state.websocket_clients:
            try:
                await client.send_text(msg)
            except Exception:
                disconnected.append(client)
        for client in disconnected:
            state.websocket_clients.remove(client)

    # Start profiling with progress polling
    async def _profile_with_progress():
        """Wrapper that runs profiling and sends periodic progress updates."""
        progress_state = {"done": 0, "total": 1}

        def on_progress(done: int, total: int):
            progress_state["done"] = done
            progress_state["total"] = total

        # Start the profiling in a thread
        import concurrent.futures
        loop = asyncio.get_event_loop()
        with concurrent.futures.ThreadPoolExecutor(max_workers=1) as pool:
            future = loop.run_in_executor(
                pool, profile_large_file, file_path, file_type,
                state.project_folder, on_progress,
            )

            # Poll and send progress while profiling runs
            while not future.done():
                await asyncio.sleep(0.5)
                msg = json.dumps({
                    "type": "profile_progress",
                    "filePath": request.filePath,
                    "done": progress_state["done"],
                    "total": progress_state["total"],
                })
                disconnected = []
                for client in state.websocket_clients:
                    try:
                        await client.send_text(msg)
                    except Exception:
                        disconnected.append(client)
                for client in disconnected:
                    state.websocket_clients.remove(client)

            # Get result (may raise if profiling failed)
            result = await future

        # Send completion
        msg = json.dumps({
            "type": "profile_complete",
            "filePath": request.filePath,
            "profile": result,
        })
        disconnected = []
        for client in state.websocket_clients:
            try:
                await client.send_text(msg)
            except Exception:
                disconnected.append(client)
        for client in disconnected:
            state.websocket_clients.remove(client)

    _profiling_task = asyncio.create_task(_profile_with_progress())

    return {"status": "profiling", "message": "Profiling started. Progress will be sent via WebSocket."}


@app.get("/api/dataframe/profile/result")
async def get_profile_result(filePath: str):
    """Get the cached profile for a file."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / filePath
    profile_path = get_profile_cache_path(state.project_folder, file_path)

    if not is_profile_valid(profile_path, file_path):
        raise HTTPException(status_code=404, detail="No valid profile found. Run profiling first.")

    profile = read_cached_profile(profile_path)
    return {"profile": profile}


@app.post("/api/dataframe/estimate-rows")
async def estimate_rows(request: EstimateRequest):
    """Estimate row count after applying filters. Fast on Parquet via predicate pushdown."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.filePath
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    ext = file_path.suffix.lower()
    file_type = "csv" if ext == ".csv" else "parquet"
    separator = _detect_csv_separator(file_path) if file_type == "csv" else ","

    try:
        count = await asyncio.to_thread(
            estimate_filtered_rows, file_path, file_type, request.filters, separator
        )
        return {"estimatedRows": count}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Estimation failed: {e}")


@app.post("/api/dataframe/filtered-preview")
async def filtered_preview(request: FilteredPreviewRequest):
    """Load a filtered subset of a massive file for preview.
    Sets up df_state so subsequent /api/dataframe/rows calls work normally."""
    if not state.project_folder:
        raise HTTPException(status_code=400, detail="No project folder selected")

    file_path = state.project_folder / request.filePath
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="File not found")

    ext = file_path.suffix.lower()

    # Set up df_state
    df_state.clear()
    df_state.file_path = file_path.as_posix()
    if ext == ".csv":
        df_state.file_type = "csv"
        df_state.csv_separator = _detect_csv_separator(file_path)
    elif ext in {".parquet", ".geoparquet"}:
        df_state.file_type = "parquet"
    else:
        raise HTTPException(status_code=400, detail="Unsupported file type")

    lf = df_state._get_lazy_frame()
    if lf is None:
        raise HTTPException(status_code=500, detail="Failed to open file")

    df_state.columns = lf.collect_schema().names()
    schema = lf.collect_schema()

    # Apply the user's filters
    df_state.current_filters = request.filters
    df_state.invalidate_filter_cache()

    # Get total filtered rows
    # Apply row limit if specified
    df_state.row_limit = request.rowLimit if request.rowLimit else None
    df_state.invalidate_filter_cache()

    filtered_lf = df_state._apply_filters_sort(lf)
    if df_state.row_limit:
        filtered_lf = filtered_lf.slice(0, df_state.row_limit)
    total_filtered = filtered_lf.select(pl.len()).collect().item()
    df_state.total_rows = total_filtered

    # Compute column info on the filtered lazy frame
    try:
        column_info = _compute_column_info(filtered_lf, df_state.columns, schema)
    except Exception:
        column_info = {col: {"type": "categorical", "values": [], "count": 0, "nullCount": 0, "blankCount": 0, "uniqueCount": 0} for col in df_state.columns}
    df_state.column_info = column_info

    # Get first chunk
    CHUNK_SIZE = 200
    rows, total_rows = df_state.get_rows(0, CHUNK_SIZE)

    print(f"[Filtered Preview] {file_path.name}: {total_filtered} rows after filters")

    return {
        "type": "dataframe",
        "filePath": request.filePath,
        "columns": df_state.columns,
        "columnInfo": column_info,
        "data": rows,
        "totalRows": total_rows,
        "offset": 0,
        "limit": CHUNK_SIZE,
        "filename": file_path.name,
    }


# WebSocket for real-time updates

@app.websocket("/ws/watch")
async def websocket_watch(websocket: WebSocket):
    """WebSocket for file change notifications"""
    await websocket.accept()
    state.websocket_clients.append(websocket)

    try:
        while True:
            # Keep connection alive, wait for messages
            try:
                data = await asyncio.wait_for(websocket.receive_text(), timeout=30)
                # Handle any incoming messages (e.g., ping)
                if data == "ping":
                    await websocket.send_text("pong")
            except asyncio.TimeoutError:
                # Send keepalive
                await websocket.send_text('{"type": "keepalive"}')
    except WebSocketDisconnect:
        state.websocket_clients.remove(websocket)
    except Exception:
        if websocket in state.websocket_clients:
            state.websocket_clients.remove(websocket)


async def notify_data_change():
    """Notify all WebSocket clients of data change"""
    if state.project_folder:
        await asyncio.to_thread(generate_metadata, state.project_folder)

    message = '{"type": "data_change"}'
    disconnected = []

    for client in state.websocket_clients:
        try:
            await client.send_text(message)
        except Exception:
            disconnected.append(client)

    for client in disconnected:
        state.websocket_clients.remove(client)


async def notify_script_change(script_path: Path):
    """Notify all WebSocket clients of script change"""
    # Send full absolute path (same format as /api/scripts endpoint)
    full_path = str(script_path)
    # Use forward slashes for consistency on Windows
    full_path = full_path.replace("\\", "/")

    # Debounce: skip if we notified about this script in the last 3 seconds
    # Use lowercase key for case-insensitive matching (Windows paths)
    debounce_key = full_path.lower()
    now = time.time()
    if debounce_key in state.last_script_change:
        if now - state.last_script_change[debounce_key] < 3.0:
            print(f"[Script Change] Debounced (duplicate within 3s): {full_path}")
            return
    state.last_script_change[debounce_key] = now
    # Clean up old entries
    state.last_script_change = {k: v for k, v in state.last_script_change.items() if now - v < 10.0}

    print(f"[Script Change] Notifying {len(state.websocket_clients)} clients: {full_path}")
    message = json.dumps({"type": "script_change", "path": full_path})
    disconnected = []

    for client in state.websocket_clients:
        try:
            await client.send_text(message)
        except Exception:
            disconnected.append(client)

    for client in disconnected:
        state.websocket_clients.remove(client)


async def notify_output_file_change(file_path: Path, change_type: str):
    """Notify all WebSocket clients of output file change for auto-preview"""
    # Get relative path from project folder
    rel_path = str(file_path)
    if state.project_folder:
        try:
            rel_path = str(file_path.relative_to(state.project_folder))
        except ValueError:
            pass
    # Use forward slashes for consistency (Windows fix)
    rel_path = rel_path.replace("\\", "/")

    print(f"[Output Change] Notifying {len(state.websocket_clients)} clients: {rel_path}")
    message = json.dumps({"type": "output_file_change", "path": rel_path, "change_type": change_type})
    disconnected = []

    for client in state.websocket_clients:
        try:
            await client.send_text(message)
        except Exception:
            disconnected.append(client)

    for client in disconnected:
        state.websocket_clients.remove(client)


# Local Terminal WebSocket

def set_terminal_size(fd, rows, cols):
    """Set terminal window size"""
    if sys.platform == 'win32':
        return  # Not supported on Windows
    winsize = struct.pack("HHHH", rows, cols, 0, 0)
    fcntl.ioctl(fd, termios.TIOCSWINSZ, winsize)


@app.websocket("/ws/terminal")
async def websocket_terminal(websocket: WebSocket):
    """WebSocket for local terminal"""
    await websocket.accept()

    # Terminal not supported on Windows
    if sys.platform == 'win32':
        await websocket.send_text("Terminal not supported on Windows.\r\n")
        await websocket.close()
        return

    # Fork a PTY
    pid, fd = pty.fork()

    if pid == 0:
        # Child process - create new session/process group so we can kill all children
        os.setsid()
        cwd = str(state.project_folder) if state.project_folder else str(Path.home())
        os.chdir(cwd)
        os.environ["TERM"] = "xterm-256color"
        os.execvp("bash", ["bash", "-l"])
    else:
        # Parent process - relay data
        print(f"[Terminal] Started PTY process {pid}")
        set_terminal_size(fd, 24, 80)

        # Make fd non-blocking
        flags = fcntl.fcntl(fd, fcntl.F_GETFL)
        fcntl.fcntl(fd, fcntl.F_SETFL, flags | os.O_NONBLOCK)

        try:
            while True:
                # Check for data from terminal (non-blocking)
                r, _, _ = select.select([fd], [], [], 0.05)
                if fd in r:
                    try:
                        data = os.read(fd, 8192)
                        if data:
                            await websocket.send_text(data.decode("utf-8", errors="replace"))
                    except OSError:
                        break

                # Check for data from websocket (with timeout)
                try:
                    data = await asyncio.wait_for(websocket.receive_text(), timeout=0.05)
                    if data:
                        # Check for JSON commands
                        if data.startswith('{'):
                            try:
                                msg = json.loads(data)
                                if msg.get('type') == 'resize':
                                    rows = msg.get('rows', 24)
                                    cols = msg.get('cols', 80)
                                    set_terminal_size(fd, rows, cols)
                                elif msg.get('type') == 'ping':
                                    await websocket.send_text('{"type":"pong"}')
                            except json.JSONDecodeError:
                                pass
                        else:
                            os.write(fd, data.encode("utf-8"))
                except asyncio.TimeoutError:
                    pass
                except WebSocketDisconnect:
                    print(f"[Terminal] WebSocket disconnected, cleaning up PTY {pid}")
                    break
        finally:
            # Clean up: close fd and kill the entire process group
            print(f"[Terminal] Cleaning up PTY process {pid}")
            try:
                os.close(fd)
            except OSError:
                pass

            # Kill the entire process group (bash + all child processes like claude)
            try:
                # First try SIGTERM to the process group
                os.killpg(pid, signal.SIGTERM)
            except OSError:
                # Process group might not exist, try killing just the pid
                try:
                    os.kill(pid, signal.SIGTERM)
                except OSError:
                    pass

            # Give processes a moment to terminate gracefully
            await asyncio.sleep(0.5)

            # Force kill if still running
            try:
                os.killpg(pid, signal.SIGKILL)
            except OSError:
                try:
                    os.kill(pid, signal.SIGKILL)
                except OSError:
                    pass

            # Reap zombie process
            try:
                os.waitpid(pid, os.WNOHANG)
            except OSError:
                pass

            print(f"[Terminal] PTY process {pid} cleaned up")


# Serve static files (React app)

@app.get("/")
async def serve_index():
    """Serve the React app index.html"""
    static_dir = get_static_dir()
    index_path = static_dir / "index.html"

    if not index_path.exists():
        return JSONResponse(
            status_code=503,
            content={
                "error": "Frontend not built",
                "message": "Run 'npm run build' in the frontend directory first"
            }
        )

    return FileResponse(index_path)


# Serve root-level static files (icon.svg, manifest.json, sw.js, etc.)
@app.get("/icon.svg")
async def serve_icon():
    """Serve the favicon SVG"""
    static_dir = get_static_dir()
    icon_path = static_dir / "icon.svg"
    if icon_path.exists():
        return FileResponse(icon_path, media_type="image/svg+xml")
    return JSONResponse(status_code=404, content={"error": "icon.svg not found"})


@app.get("/icon-192.png")
async def serve_icon_192():
    static_dir = get_static_dir()
    path = static_dir / "icon-192.png"
    if path.exists():
        return FileResponse(path, media_type="image/png")
    return JSONResponse(status_code=404, content={"error": "icon-192.png not found"})


@app.get("/icon-512.png")
async def serve_icon_512():
    static_dir = get_static_dir()
    path = static_dir / "icon-512.png"
    if path.exists():
        return FileResponse(path, media_type="image/png")
    return JSONResponse(status_code=404, content={"error": "icon-512.png not found"})


@app.get("/manifest.json")
async def serve_manifest():
    static_dir = get_static_dir()
    path = static_dir / "manifest.json"
    if path.exists():
        return FileResponse(path, media_type="application/manifest+json")
    return JSONResponse(status_code=404, content={"error": "manifest.json not found"})


@app.get("/sw.js")
async def serve_service_worker():
    static_dir = get_static_dir()
    path = static_dir / "sw.js"
    if path.exists():
        return FileResponse(path, media_type="application/javascript")
    return JSONResponse(status_code=404, content={"error": "sw.js not found"})


# Mount static files for assets (at module load time)
_static_dir = get_static_dir()
_assets_dir = _static_dir / "assets"
if _assets_dir.exists():
    app.mount("/assets", StaticFiles(directory=_assets_dir), name="assets")


def create_app() -> FastAPI:
    """Factory function for creating the app"""
    return app
